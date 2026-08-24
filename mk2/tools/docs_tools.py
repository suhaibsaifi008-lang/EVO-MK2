"""Document creation tools — EVO makes REAL Word files, not outlines.

docs_create(path, markdown) renders simple markdown (headings, bullets,
numbered lists, bold, paragraphs) into a .docx via python-docx. If
python-docx is unavailable it falls back to a Word-compatible HTML file
(.doc) that Microsoft Word opens natively - so the capability never dies.

Intended flow for "write me a report":
    web_search / deep_research  ->  docs_create(report.md content)
    -> real file lands in Documents (fs_tools allowlist), path reported.
"""
import re

from .. import db
from . import tool


def _safe_path(path: str):
    from ..fs_tools import _safe

    return _safe(path)


def _md_to_blocks(md: str) -> list[tuple[str, str]]:
    """Parse lightweight markdown into (kind, text) blocks."""
    blocks: list[tuple[str, str]] = []
    for raw in (md or "").replace("\r\n", "\n").split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
        elif re.match(r"^\s*[-*]\s+", line):
            blocks.append(("li", re.sub(r"^\s*[-*]\s+", "", line).strip()))
        elif re.match(r"^\s*\d+[.)]\s+", line):
            blocks.append(("ol", re.sub(r"^\s*\d+[.)]\s+", "", line).strip()))
        else:
            blocks.append(("p", line.strip()))
    return blocks


_B_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_runs(paragraph, text: str) -> None:
    """Support **bold** inside paragraphs/bullets."""
    pos = 0
    for m in _B_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _render_docx(md: str, path):
    import docx

    d = docx.Document()
    for kind, text in _md_to_blocks(md):
        if kind == "h1":
            d.add_heading(text, level=0)
        elif kind == "h2":
            d.add_heading(text, level=1)
        elif kind == "h3":
            d.add_heading(text, level=2)
        elif kind == "li":
            _add_runs(d.add_paragraph(style="List Bullet"), text)
        elif kind == "ol":
            _add_runs(d.add_paragraph(style="List Number"), text)
        else:
            _add_runs(d.add_paragraph(), text)
    d.save(str(path))


def _render_html_doc(md: str, path):
    """python-docx missing: Word-compatible HTML saved as .doc."""
    parts = ["<html><head><meta charset='utf-8'></head><body>"]
    for kind, text in _md_to_blocks(md):
        safe = text.replace("&", "&amp;").replace("<", "&lt;")
        safe = _B_RE.sub(r"<b>\1</b>", safe)
        tag = {"h1": "h1", "h2": "h2", "h3": "h3",
               "li": "li", "ol": "li", "p": "p"}[kind]
        if kind in ("li", "ol"):
            if parts[-1] != "<ul>":
                parts.append("<ul>")
            parts.append(f"<li>{safe}</li>")
        else:
            if parts[-1] == "<ul>":
                parts.append("</ul>")
            parts.append(f"<{tag}>{safe}</{tag}>")
    parts.append("</body></html>")
    html = "\n".join(p for p in parts if p is not None)
    path = path.with_suffix(".doc")
    path.write_text(html, encoding="utf-8")
    return path


@tool("docs_create", "Create a real Word document (.docx) at a path from markdown content (# headings, - bullets, **bold**). Use this to actually DELIVER reports/files.",
      {"path": {"type": "string"}, "content": {"type": "string"}},
      permission="execute", long_running=False)
def docs_create(path: str, content: str) -> dict:
    p = _safe_path(path)
    if p.suffix.lower() not in (".docx", ".doc"):
        p = p.with_suffix(".docx")
    p.parent.mkdir(parents=True, exist_ok=True)
    md = (content or "").strip()
    if len(md) < 20:
        return {"ok": False,
                "speech": "The document content came back empty.", "data": {}}
    try:
        _render_docx(md, p)
    except ImportError:
        p = _render_html_doc(md, p)
    except Exception as exc:
        return {"ok": False, "speech": f"Document build failed: {str(exc)[:150]}",
                "data": {}}
    db.audit("docs_create", str(path)[:200], True, str(p.name))
    size_kb = max(0.1, p.stat().st_size / 1024)
    words = len(md.split())
    return {"ok": True,
            "speech": (f"Created {p.name} ({words} words, "
                       f"{size_kb:.0f} KB) at {p.parent}."),
            "data": {"path": str(p), "words": words}}


@tool("docs_append", "Append more markdown sections to an existing .docx you created earlier.",
      {"path": {"type": "string"}, "content": {"type": "string"}},
      permission="execute")
def docs_append(path: str, content: str) -> dict:
    p = _safe_path(path)
    if not p.exists():
        # chain to create when the file doesn't exist yet
        return docs_create.__wrapped__(path, content) if hasattr(
            docs_create, "__wrapped__") else docs_create(path, content)
    try:
        import docx

        d = docx.Document(str(p))
        for kind, text in _md_to_blocks(content):
            if kind == "h1":
                d.add_heading(text, level=0)
            elif kind == "h2":
                d.add_heading(text, level=1)
            elif kind == "h3":
                d.add_heading(text, level=2)
            elif kind == "li":
                _add_runs(d.add_paragraph(style="List Bullet"), text)
            elif kind == "ol":
                _add_runs(d.add_paragraph(style="List Number"), text)
            else:
                _add_runs(d.add_paragraph(), text)
        d.save(str(p))
    except ImportError:
        return {"ok": False,
                "speech": "Appending needs python-docx (pip install python-docx).",
                "data": {}}
    except Exception as exc:
        return {"ok": False, "speech": f"Append failed: {str(exc)[:150]}", "data": {}}
    db.audit("docs_append", str(path)[:200], True, p.name)
    return {"ok": True, "speech": f"Added more sections to {p.name}.",
            "data": {"path": str(p)}}


def read_back(path: str) -> str:
    """Test helper: extract all text from a generated docx."""
    import docx

    d = docx.Document(str(path))
    return "\n".join(par.text for par in d.paragraphs)
